#!/usr/bin/env python3
"""生成实验报告 Word 文档 — 嵌入所有图片 + 回答所有思考题"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = r"C:\Users\Administrator\Desktop\5"
FONT_NAME = '微软雅黑'

doc = Document()

def set_cjk_font(run, font_name=FONT_NAME, size=None, bold=False, color=None, italic=False):
    """在每个 run 上同时设置西文和东亚字体"""
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    # 移除已有的 rFonts 再插入新的
    existing = rPr.findall(qn('w:rFonts'))
    for e in existing:
        rPr.remove(e)
    rPr.insert(0, rFonts)
    run.font.name = font_name
    if size:
        run.font.size = size
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    if italic:
        run.font.italic = True

def add_para(text, bold=False, size=None, color=None, alignment=None):
    """添加段落（自动处理 CJK 字体）"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    set_cjk_font(run, size=size or Pt(11), bold=bold, color=color)
    return p

def add_heading_text(text, level=1):
    """添加标题（绕过样式系统，手动设置格式）"""
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(text)
        set_cjk_font(run, size=Pt(16), bold=True, color=RGBColor(0x3C, 0x54, 0x88))
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run = p.add_run(text)
        set_cjk_font(run, size=Pt(14), bold=True, color=RGBColor(0x3C, 0x54, 0x88))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run = p.add_run(text)
        set_cjk_font(run, size=Pt(12), bold=True, color=RGBColor(0x3C, 0x54, 0x88))
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
    return p

def add_bullet(text):
    """添加项目符号段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('• ' + text)
    set_cjk_font(run, size=Pt(11))
    return p

def add_image(img_name, caption, width=5.5):
    """插入图片和题注"""
    img_path = os.path.join(BASE_DIR, img_name)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        set_cjk_font(cap_run, size=Pt(9), italic=True, color=RGBColor(0x84, 0x91, 0xB4))

def add_code(code_text):
    """添加代码块（灰底等宽字体）"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        set_cjk_font(run, font_name='Consolas', size=Pt(8), color=RGBColor(0x33, 0x33, 0x33))

def add_table(headers, rows):
    """添加表格（自动处理 CJK 字体）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_cjk_font(run, size=Pt(10), bold=True)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_cjk_font(run, size=Pt(10))
    doc.add_paragraph()

# ═══════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════
doc.add_paragraph()
add_para('第五部分：LLM 与大模型应用 —— 农业知识问答系统',
         bold=True, size=Pt(22), color=RGBColor(0x3C, 0x54, 0x88),
         alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('实 验 报 告', size=Pt(16),
         color=RGBColor(0x84, 0x91, 0xB4),
         alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# ═════════════ 实验环境 ═════════════
add_heading_text('实验环境', level=1)
add_table(
    ['配置项', '详情'],
    [
        ['本地环境', 'Windows 11 + Python 3.12 + Docker (dl-env)'],
        ['云端 LLM API', 'agicto API (gpt-4o-mini)'],
        ['向量模型', 'BAAI/bge-small-zh-v1.5 (512维)'],
        ['知识库', 'ChromaDB 持久化存储'],
        ['GPU 服务器', 'Tesla V100S 32GB, CUDA 13.0'],
        ['本地部署模型', 'Qwen2.5-7B-Instruct (FP16, ~15GB 显存)'],
        ['推理框架', 'Flask + HuggingFace Transformers + PyTorch 2.6.0'],
        ['可视化', 'Matplotlib (NPG Nature Publishing Group 顶刊配色)'],
        ['Web 框架', 'Streamlit'],
    ]
)

# ═════════════ 任务 A ═════════════
add_heading_text('任务 A：LLM API 初体验与 Prompt Engineering', level=1)

add_heading_text('实验目的', level=2)
add_para('调用 agicto API 进行对话，对比三种 Prompt 策略在农业场景下的效果。')

add_heading_text('实验设计', level=2)
add_table(
    ['策略', 'System Prompt', '特点'],
    [
        ['基础 Prompt', '无', '直接提问，无额外引导'],
        ['角色设定 + 结构化', '20年农业植保专家', '指定身份 + 要求结构化输出'],
        ['Few-shot + Chain-of-Thought', '专家 + 诊断示例', '2个示例 + 逐步推理引导'],
    ]
)

add_heading_text('实验结果', level=2)
add_para('测试问题：番茄叶子发黄是什么原因？')
add_table(
    ['策略', '响应时间', '回答长度', '质量评分'],
    [
        ['基础 Prompt', '6.2s', '~500 chars', '3/10'],
        ['角色设定 + 结构化', '7.3s', '~700 chars', '7/10'],
        ['Few-shot + CoT', '3.2s', '~350 chars', '9/10'],
    ]
)

add_heading_text('可视化结果', level=2)
add_image('task_a_prompt_comparison.png', '图1: 三种 Prompt 策略的质量评分对比。Few-shot + Chain-of-Thought 策略获得最高评分（9/10）。')
add_image('task_a_response_time.png', '图2: 各策略的响应时间与回答长度双轴对比。')
add_image('task_a_radar.png', '图3: 多维评估雷达图。Few-shot 策略在专业性、结构性、准确性等维度全面领先。')

add_heading_text('多轮对话测试', level=2)
add_para('模拟农业咨询场景的三轮对话（病害诊断 → 用药咨询 → 施用频率），模型展现出良好的上下文理解能力。')

add_heading_text('结论', level=2)
add_bullet('角色设定显著提升回答的专业性和可操作性')
add_bullet('Few-shot + CoT 策略在专业诊断场景中效果最佳')
add_bullet('结构化输出要求对农业技术指导场景至关重要')

# ═════════════ 任务 B ═════════════
add_heading_text('任务 B：RAG 农业知识库问答系统', level=1)

add_heading_text('实验目的', level=2)
add_para('构建基于检索增强生成（RAG）的农业知识问答系统，对比有/无知识库的回答质量。')

add_heading_text('系统架构', level=2)
add_code('知识文档 → 文本分块(500字/块, 50字重叠) → BGE向量化(512维) → ChromaDB\n'
         '                                                      ↓\n'
         '用户提问 → 向量化 → 检索Top-K相关块 → 拼接Prompt → LLM生成回答')

add_heading_text('知识库构成', level=2)
add_image('task_b_kb_composition.png', '图4: 知识库包含4篇文档（diseases/planting/pesticide），共18个文本块。')

add_heading_text('检索效果', level=2)
add_para('对三个典型问题的检索结果：')
add_image('task_b_retrieval_relevance.png', '图5: 三个问题的检索相关度。Top-3 检索的平均相关度在 0.70-0.80 之间。')
add_image('task_b_heatmap.png', '图6: 检索相关度热力图。同类文档检索效果最佳。')

add_heading_text('RAG vs Direct LLM 对比', level=2)
add_image('task_b_rag_comparison.png', '图7: 有 RAG 的回答能提供具体药品名称、浓度和安全间隔期，而无 RAG 的回答较为笼统。')

add_heading_text('关键发现', level=2)
add_bullet('减少幻觉：RAG 将回答锚定在实际文档上，大幅减少编造信息')
add_bullet('专业深度：知识库中的专业内容（农药浓度、安全间隔期）直接被引用')
add_bullet('可追溯性：每个回答都能追溯到具体的参考文档')

# ═════════════ 任务 C ═════════════
add_heading_text('任务 C：GPU 服务器本地模型部署与性能对比', level=1)

add_heading_text('实验目的', level=2)
add_para('在 GPU 服务器上部署开源模型，对比本地部署与云端 API 的性能差异。')

add_heading_text('部署配置', level=2)
add_table(
    ['项目', '配置'],
    [
        ['GPU', 'Tesla V100S 32GB'],
        ['模型', 'Qwen2.5-7B-Instruct (FP16, ~14GB 显存)'],
        ['推理框架', 'Flask + HuggingFace Transformers'],
        ['API 接口', 'OpenAI 兼容 (/v1/chat/completions)'],
    ]
)

add_heading_text('性能对比', level=2)
add_table(
    ['指标', '本地 GPU (V100S)', '云端 API (gpt-4o-mini)'],
    [
        ['平均响应时间', '20.11s', '5.98s'],
        ['平均输出长度', '817 chars', '659 chars'],
        ['字符生成速度', '40.6 chars/s', '110.2 chars/s'],
        ['显存占用', '~15 GB', 'N/A'],
    ]
)

add_heading_text('可视化结果', level=2)
add_image('task_c_response_time.png', '图8: 本地 GPU vs 云端 API 响应时间对比。云端 API 响应速度快 3.4 倍。')
add_image('task_c_dual_comparison.png', '图9: 响应时间与输出长度双轴对比。本地模型输出更详细（+24%），但速度较慢。')
add_image('task_c_throughput.png', '图10: 字符生成速度对比。云端 API 的吞吐量是本地 V100S 的 2.7 倍。')

add_heading_text('讨论', level=2)
add_bullet('云端优势：gpt-4o-mini 在速度上明显领先（5.98s vs 20.11s），得益于更强的基础设施')
add_bullet('本地优势：无 API 费用、数据不出服务器、可完全控制模型行为')
add_bullet('V100S 局限性：较老架构，FP16 推理速度有限；使用 vLLM 或 TensorRT-LLM 可显著提升')

# ═════════════ 任务 D ═════════════
add_heading_text('任务 D：农业 AI 助手 Web 应用', level=1)

add_heading_text('功能特性', level=2)
add_bullet('Streamlit Web UI：简洁的聊天界面，侧边栏配置 API Key 和模型')
add_bullet('RAG 引擎：BGE-small-zh-v1.5 向量检索 + ChromaDB 持久化知识库')
add_bullet('多模型支持：可切换 gpt-4o-mini / qwen-plus')
add_bullet('参考来源展示：每次回答附带检索到的文档来源和相关度')

add_heading_text('运行方式', level=2)
add_code('streamlit run task_d_challenge.py --server.address 0.0.0.0 --server.port 8501')
add_para('应用地址：http://localhost:8501')

add_heading_text('系统截图', level=2)
add_image('task_d_streamlit_app.png',
    '图11: Streamlit 农业 AI 助手 Web 应用。测试问题："番茄叶片出现褐色斑点，是什么病？怎么防治？"。'
    '系统通过 RAG 检索知识库，返回了包括晚疫病和早疫病的详细诊断及具体防治方案'
    '（含农药名称、浓度、安全间隔期），并附带可追溯的参考来源。',
    width=5.8)

# ═════════════ 任务总结 ═════════════
add_heading_text('任务总结', level=1)

add_heading_text('完成情况', level=2)
add_table(
    ['任务', '状态', '关键产出'],
    [
        ['Task A: Prompt Engineering', '已完成', '3种策略对比 + 3张NPG图表'],
        ['Task B: RAG 系统', '已完成', '向量知识库 + RAG vs no-RAG 对比 + 4张图表'],
        ['Task C: 本地部署', '已完成', 'Flask + Qwen2.5-7B 部署 + 性能对比 + 3张图表'],
        ['Task D: Web 应用', '已完成', 'Streamlit 农业 AI 助手（含功能截图）'],
    ]
)

add_heading_text('关键技术栈', level=2)
add_bullet('LLM API: agicto (兼容 OpenAI SDK)')
add_bullet('向量模型: BAAI/bge-small-zh-v1.5 (512维)')
add_bullet('向量数据库: ChromaDB (持久化存储)')
add_bullet('本地推理: Flask + HuggingFace Transformers + PyTorch 2.6.0')
add_bullet('可视化: Matplotlib (NPG Nature Publishing Group 顶刊配色)')
add_bullet('Web 框架: Streamlit')

# ═══════════════════════════════════════════
# 思考题回答
# ═══════════════════════════════════════════
add_heading_text('思考题回答', level=1)

# ── Task A ──
add_heading_text('Task A 思考题', level=2)

add_heading_text('1. 为什么角色设定能显著提升 LLM 的回答质量？', level=3)
add_para('角色设定（Role Prompting）能够显著提升 LLM 回答质量，原因包括以下几个方面：')
add_para('（1）激活领域知识模式。LLM 在预训练阶段学习了海量文本，其中包含了不同身份、不同专业背景的表达方式。当给定"农业植保专家"这一角色时，模型会激活训练数据中与该角色相关的词汇、知识结构和表达风格，使输出更贴合专业场景。这种机制可以理解为一种"条件生成"——角色设定作为条件信号，约束了模型在特定知识子空间中进行采样。')
add_para('（2）缩小输出空间。通用 LLM 在面对开放问题时，可能的输出空间非常大，容易产生泛泛而谈的回答。角色设定相当于给模型划定了一个"专业边界"，缩小了输出的可能性空间，让模型聚焦于与角色相关的知识和表达方式。')
add_para('（3）提升回答的结构性和可操作性。角色设定往往隐含了该角色的沟通习惯和思维方式。例如，农业专家在面对病害咨询时，自然会按照"症状识别→病因分析→防治方案"的逻辑来组织回答，而非简单地罗列知识点。这种隐含的结构性指导使回答更加实用。')
add_para('（4）增强可信度。角色设定还可以影响回答的语气和措辞，让回答听起来更专业、更有说服力。对于需要实操指导的农业场景，这种可信度尤为重要。')
add_para('本实验中，角色设定（策略2）将质量评分从基础 Prompt 的 3/10 提升到 7/10，验证了这一结论。')

add_heading_text('2. Few-shot learning 中，示例的数量和质量哪个更重要？', level=3)
add_para('在 Few-shot learning 中，示例的质量比数量更重要，但两者存在一个最优平衡点。')
add_para('（1）质量优先的理由：LLM 对示例的模式非常敏感。高质量的示例能够精准地展示输入-输出的映射关系、推理步骤和输出格式，模型可以从中快速"学会"期望的行为模式。一个精心设计的示例比三个粗糙的示例更有价值。低质量的示例（如推理逻辑混乱、格式不一致、信息不准确）会误导模型，甚至降低回答质量。在本实验中，仅使用 2 个高质量的诊断示例（策略3），就获得了最高评分 9/10。')
add_para('（2）数量的作用：当任务模式较为复杂或多变时，增加示例数量可以帮助模型更好地理解任务的边界和多样性。但示例数量存在边际收益递减规律——从 0 到 2-3 个示例的收益最大，超过 5 个后收益急剧下降。过多的示例（尤其对较短的上下文窗口）会占用宝贵的 token 空间，反而压缩了输出空间。')
add_para('（3）实际建议：优先保证示例质量（准确性、代表性、格式一致性），在此基础上提供 2-4 个覆盖不同情况的示例。对于农业诊断场景，应选择典型、特征明确、推理链清晰的案例作为示例。')

add_heading_text('3. 在农业场景下，结构化输出为什么特别重要？', level=3)
add_para('农业场景下结构化输出特别重要，原因如下：')
add_para('（1）实操性强。农业技术指导（如病害防治）通常包含多个相互关联但逻辑不同的部分——症状识别、病因分析、农业防治、化学防治、安全间隔期等。结构化输出能将这些信息分类呈现，便于农民或技术人员快速定位所需内容。例如，一个农民可能更关心"用什么药"和"打多少"，而非病原菌的学名。')
add_para('（2）减少误用风险。农业用药涉及浓度、安全间隔期等关键安全参数。如果这些信息混杂在冗长的段落中，用户可能忽略或记错，导致农药残留超标或药害。结构化输出（如表格形式列出药剂、浓度、间隔期）能显著降低此类风险。')
add_para('（3）便于标准化推广。农业技术推广需要标准化的信息传递格式。结构化输出便于制作技术手册、培训材料和数字化的农业知识库，实现知识的标准化管理和传播。')
add_para('（4）提升用户信任。结构清晰、信息完整的回答给用户以"专业可靠"的印象，增强用户对 AI 系统的信任度。')
add_para('本实验中，策略2（角色设定+结构化）相比策略1（基础Prompt）评分从 3/10 跃升至 7/10，结构化要求在其中的贡献不可忽视。')

# ── Task B ──
add_heading_text('Task B 思考题', level=2)

add_heading_text('1. RAG 系统为什么能减少 LLM 的"幻觉"问题？', level=3)
add_para('RAG（检索增强生成）能有效减少 LLM 幻觉，其核心机制在于将生成范式从"自由回忆"转变为"阅读理解"：')
add_para('（1）锚定效应。RAG 将检索到的真实文档片段作为 LLM 生成的"锚点"。模型不再是凭训练记忆凭空生成内容，而是基于给定的参考资料进行"受约束的生成"。这大大降低了模型编造不存在信息的概率。')
add_para('（2）认知负荷转移。在没有 RAG 的情况下，LLM 需要同时完成"回忆知识"和"组织语言"两个任务。训练数据中农业领域的细节（如特定农药的浓度、安全间隔期）可能记忆不准确或被遗忘。RAG 将"知识供给"外包给检索系统，让 LLM 专注于"语言组织"，从而减少因记忆不准产生的幻觉。')
add_para('（3）可验证性。RAG 的每个回答都可以追溯到具体的参考文档片段。如果回答有误，可以定位到是检索阶段（检索到了不相关的内容）还是生成阶段（LLM 误解了检索内容）出了问题。而无 RAG 的回答无法验证其出处。')
add_para('（4）温度参数的配合。在 RAG 场景中，可以将 temperature 设得更低（如 0.3），因为创意性不再重要，准确性才是关键。低温度进一步减少了模型的随机采样，降低了幻觉产生的概率。')
add_para('本实验中，无 RAG 的回答较为笼统（如"使用杀菌剂"），而有 RAG 的回答给出了具体的药品名称、浓度和安全间隔期，这些具体信息直接来源于知识库文档。')

add_heading_text('2. 文本分块的大小（chunk_size）对检索效果有什么影响？太大或太小会怎样？', level=3)
add_para('文本分块大小是 RAG 系统中关键的超参数，直接影响检索质量和生成效果：')
add_para('（1）chunk_size 过小（如 <200 字）：优点是检索精度高，能精准匹配到特定细节（如某个农药的浓度值）。缺点是语义信息不完整，一个完整的知识点（如"早疫病的完整防治方案"）可能被切分到多个块中，导致检索到的上下文碎片化，LLM 无法获得完整信息来生成连贯的回答。极端情况下，每个句子一个 chunk，检索变成了关键词匹配而非语义匹配。')
add_para('（2）chunk_size 过大（如 >2000 字）：优点是包含完整上下文，LLM 可以看到一个主题的全貌。缺点是检索精度下降，chunk 中可能包含大量与查询无关的内容（噪音），稀释了关键信息的密度。同时，向量嵌入对长文本的语义表示能力也会下降——长文本的向量是全文语义的"平均"，特定细节可能被淹没。')
add_para('（3）最优策略：中等大小（300-800 字）+ 适当重叠（10-20%）是大多数场景的最佳实践。本实验采用 500 字/块 + 50 字重叠的配置，在语义完整性和检索精度之间取得了较好的平衡。对于包含表格（如农药浓度表）的结构化内容，应考虑按表格的语义单元（如一个完整的表格行组）进行分块。')
add_para('（4）进阶思路："多粒度检索"——同时维护不同 chunk_size 的索引，先用小块检索精确信息，再用大块补充上下文，最后合并给 LLM。')

add_heading_text('3. 如果知识库中没有相关内容，系统应该如何处理？', level=3)
add_para('这是一个非常重要的工程问题。RAG 系统必须优雅地处理知识库覆盖不足的情况，否则 LLM 可能会基于不相关的检索结果强行"编造"答案，反而比不使用 RAG 更差。推荐以下多层级的处理策略：')
add_para('（1）检索质量阈值过滤。为检索结果设置一个相关度阈值（如 cosine similarity > 0.6 或 distance < 0.4）。如果 Top-K 结果的相关度都低于阈值，说明知识库中没有相关内容，此时应触发"知识不足"分支。')
add_para('（2）诚实告知 + 降级回答。当知识库覆盖不足时，系统应明确告知用户"抱歉，当前知识库中没有关于该问题的直接资料"，同时可以选择性地提供通用建议（但需标注"以下信息未经验证"）。')
add_para('（3）降级为直接 LLM 回答（谨慎使用）。作为备选方案，可以在无相关检索结果时回退到直接 LLM 回答，但必须明确标注"以下回答未基于知识库，仅供参考"。')
add_para('（4）记录未命中查询。将知识库未能覆盖的查询记录下来，作为后续扩充知识库的依据。这是 RAG 系统持续改进的重要数据来源。')
add_para('（5）本实验中的处理。系统 Prompt 中明确要求"如果参考资料不足以回答问题，请如实告知"，从生成侧约束 LLM 不要强行编造。这是最简单但有效的兜底策略。最佳实践是结合检索侧（阈值过滤）和生成侧（Prompt 约束）的双重保障。')

add_heading_text('4. 中文向量模型和英文模型有什么区别？为什么要用专门的中文模型？', level=3)
add_para('中文向量模型和英文模型在多个层面存在显著差异，使用专门的中文模型对于中文 RAG 系统至关重要：')
add_para('（1）分词粒度不同。英文以空格分词，单词是自然的语义单元。英文模型的 tokenizer 基于 BPE（Byte Pair Encoding），对英文单词的切分较为自然。中文没有天然的分词边界，同一个句子可以有多种切分方式（如"农业知识库"可切为"农业/知识库"或"农/业/知识/库"）。中文模型（如 BGE-small-zh）的 tokenizer 专门针对中文语料训练，能更合理地处理中文分词。')
add_para('（2）语义表示能力不同。通用多语言模型（如 multilingual-e5）虽然支持中文，但在中文语义理解上的精度通常不如专门的中文模型。因为多语言模型的训练数据需要在多种语言之间分配，中文数据的比例和深度有限。专门的中文模型（如 BAAI/bge-small-zh-v1.5）在海量中文语料上训练，对中文特有的表达方式、成语、专业术语有更好的理解。')
add_para('（3）向量维度与效率。英文/多语言模型通常维度较高（如 768 或 1024 维），计算和存储开销更大。中文模型可以在较低维度（如 512 维）下保持良好的语义区分能力，更高效。本实验使用的 bge-small-zh-v1.5 仅 512 维，在农业知识库上检索相关度达到 0.70-0.80，效果良好。')
add_para('（4）农业领域的特殊性。农业中文文本含有大量专业术语（如"代森锰锌""烯酰吗啉""同心轮纹"等），这些术语在英文或多语言模型中可能被错误切分或无法准确理解其语义关联。专门的中文模型能更好地处理这些专业表达。')
add_para('综上，对于中文 RAG 系统，使用专门的中文向量模型在检索精度、计算效率和领域适应方面都优于通用多语言模型。')

# ── Task C ──
add_heading_text('Task C 思考题', level=2)

add_heading_text('1. vLLM 相比直接用 HuggingFace Transformers 推理有什么优势？', level=3)
add_para('vLLM 是专为高性能 LLM 推理设计的框架，相比直接使用 HuggingFace Transformers 有以下核心优势：')
add_para('（1）PagedAttention 技术。这是 vLLM 最核心的创新。传统 HuggingFace 推理中，KV Cache 的显存管理采用连续分配方式，会产生大量显存碎片（类似操作系统中的内存碎片问题），实际显存利用率仅 20-40%。PagedAttention 借鉴操作系统的虚拟内存分页机制，将 KV Cache 切分为固定大小的"页"，按需分配和回收，显存利用率可提升至接近 100%，几乎消除浪费。')
add_para('（2）连续批处理（Continuous Batching）。传统推理每次处理一个请求，新请求需要等当前批次完成。vLLM 支持在生成过程中动态插入新请求，当一个序列生成结束后立即释放资源并接收新请求，大幅提升并发吞吐量。官方数据显示吞吐量可达 HuggingFace 的 24 倍。')
add_para('（3）高效的内核实现。vLLM 使用 CUDA 级别的优化内核（如 FlashAttention），在 GPU 计算层面减少显存读写，进一步加速推理。这些优化对普通用户是透明的，无需手动配置。')
add_para('（4）OpenAI 兼容 API。vLLM 原生提供与 OpenAI API 兼容的 HTTP 服务接口（/v1/chat/completions），部署后可直接用 OpenAI SDK 调用，降低了系统集成成本。')
add_para('（5）生产级特性。支持量化（AWQ、GPTQ）、多 GPU 张量并行、前缀缓存（Prefix Caching）等生产环境必需的功能。')
add_para('注：本实验中，受限于 V100S 的 CUDA 兼容性问题（vLLM 0.21.0+ 对较老架构支持有限），最终使用了 Flask + HuggingFace Transformers 的方式部署。但在较新的 GPU（如 A100、H100）上，vLLM 是更优选择。')

add_heading_text('2. 如果显存不够，有哪些方法可以降低显存占用？', level=3)
add_para('显存不足是本地部署 LLM 时最常见的问题。以下是降低显存占用的主要方法，按效果从大到小排列：')
add_para('（1）模型量化（Quantization）。INT8 量化将模型参数从 FP16（每个参数 2 字节）压缩到 INT8（每个参数 1 字节），显存需求减半，精度损失通常 < 1%。INT4 量化（GPTQ/AWQ）将参数压缩到 4 位（每个参数 0.5 字节），显存需求降至 FP16 的 1/4。例如 Qwen2.5-7B 从约 14GB 降至约 4-5GB，精度损失约 2-5%，但在多数应用场景下仍可接受。GGUF 格式（llama.cpp）专为 CPU/边缘设备设计，支持 2-8 位量化。')
add_para('（2）选择更小的模型。7B → 3B → 1.5B 参数量的阶梯式降级。虽然能力会下降，但在特定垂直领域（如农业），通过微调可以弥补参数量的不足。例如 Qwen2.5-1.5B-Instruct 仅需约 3GB 显存（FP16），在简单问答场景下仍有不错表现。')
add_para('（3）减少上下文长度（max-model-len）。KV Cache 的显存占用与序列长度成正比。将 max-model-len 从 8192 降至 4096 或 2048，可以节省数 GB 显存。')
add_para('（4）降低批处理大小（batch size）。同时处理的请求数越少，KV Cache 占用越小。在低并发场景下可以显著节省显存。')
add_para('（5）使用 FlashAttention。将注意力计算从 O(n²) 显存降为 O(n)，对大上下文场景节省明显。vLLM 和较新版本的 Transformers 已默认启用。')
add_para('（6）CPU Offloading。将部分层或 KV Cache 卸载到 CPU 内存，牺牲速度换取更大的可用空间。')
add_para('本实验中，Qwen2.5-7B (FP16) 占用约 15GB，在 32GB 的 V100S 上运行绰绰有余。如果使用 24GB 的 RTX 3090，FP16 模式也可正常运行（剩余约 9GB 给 KV Cache）。')

add_heading_text('3. 本地部署和云端 API 各有什么优劣？在实际项目中如何选择？', level=3)
add_para('本地部署和云端 API 各有优劣势，选择需要综合考虑多个维度。')
add_para('一、云端 API 的优势与劣势：')
add_para('优势：（1）零运维成本，无需购买和维护 GPU 硬件，无需处理驱动、CUDA 版本兼容等问题。（2）弹性伸缩，按需付费，可根据流量自动扩缩容，适合流量波动大的场景。（3）高性能，云厂商的 GPU 集群（A100/H100）通常比个人/小团队能获得的硬件更强大。本实验中云端 gpt-4o-mini 的响应速度（5.98s）是本地 V100S（20.11s）的 3.4 倍。（4）快速迭代，无需等待模型下载和部署，API 调用即可开始开发。')
add_para('劣势：（1）数据隐私风险，敏感数据（如农业企业的核心种植数据）需要发送到云端，存在合规和隐私隐患。（2）长期成本高，高频调用场景下，API 费用可能远超购买 GPU 的一次性成本。（3）可控性差，无法自定义模型行为、无法微调、受限于 API 提供商的速率限制和服务可用性。（4）网络依赖，离线或网络不稳定的农业场景（如田间地头）无法使用。')
add_para('二、本地部署的优势与劣势：')
add_para('优势：（1）数据安全，全部数据留在本地服务器，适合处理敏感农业数据或企业机密。（2）无调用成本，高频调用场景下边际成本接近零，适合大规模农业知识服务。（3）完全可控，可自由选择模型、微调、量化策略，可根据农业领域定制优化。（4）离线可用，不依赖网络连接，适合偏远地区的农业应用场景。')
add_para('劣势：（1）硬件投入高，GPU 服务器一次性投入大（V100S 约 3-5 万元，A100 约 8-15 万元）。（2）运维维护，需要维护硬件、更新驱动、监控服务状态。（3）性能受限，个人/小团队通常只能获得 1-2 块 GPU，并发能力和推理速度有限。本实验中本地 V100S 的字符生成速度（40.6 chars/s）仅为云端（110.2 chars/s）的 37%。')
add_para('三、实际项目选择建议：')
add_bullet('原型验证 / 低频调用 → 云端 API（快速启动、成本低）')
add_bullet('数据敏感的农业企业 → 本地部署（数据不出门）')
add_bullet('高并发 / 大规模服务 → 需权衡：云端弹性 + 便捷 vs 本地低成本 + 可控')
add_bullet('推荐混合架构：日常使用云端 API，核心敏感数据用本地模型处理，两者通过统一接口调用')
add_bullet('离线 / 边缘场景（如田间诊断） → 本地部署轻量化模型（INT4 量化 + 小参数模型）')
add_para('本实验展示了两种方案的完整实现：云端 agicto API (Task A/B/D) 和本地 Flask + Qwen2.5-7B (Task C)，为实际项目中的架构选择提供了参考依据。')

# ── 保存 ──
output_path = os.path.join(BASE_DIR, '农业AI实验报告_终稿_v2.docx')
doc.save(output_path)
print(f'Word文档已生成: {output_path}')
